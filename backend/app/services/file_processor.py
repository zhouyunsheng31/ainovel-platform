"""
AI小说拆书系统 - 文件处理服务
支持 epub/txt/doc/docx/pdf 格式解析
"""
import re
import chardet
from typing import Tuple
from pathlib import Path

from app.config import settings


class FileProcessor:
    """文件处理服务 - 解码检测与文本提取"""
    
    SUPPORTED_FORMATS = ['.txt', '.epub', '.doc', '.docx', '.pdf']
    
    # 编码检测优先级
    ENCODING_PRIORITY = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'utf-16']
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def save_file(self, file_content: bytes, filename: str) -> Tuple[str, str]:
        """
        保存上传文件
        返回: (文件路径, 文件类型)
        """
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        # 生成唯一文件名
        import uuid
        unique_name = f"{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / unique_name
        
        # 异步写入
        import aiofiles
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        return str(file_path), file_ext[1:].upper()  # 去掉点，转大写
    
    def detect_encoding(self, file_content: bytes) -> str:
        """
        检测文本编码
        优先尝试 UTF-8，失败则使用 chardet
        """
        # 1. 优先尝试 UTF-8
        for encoding in self.ENCODING_PRIORITY:
            try:
                file_content.decode(encoding)
                return encoding
            except (UnicodeDecodeError, LookupError):
                continue
        
        # 2. 使用 chardet 检测
        result = chardet.detect(file_content)
        detected = result.get('encoding', 'utf-8')
        
        # chardet 可能返回 None 或不存在的编码
        if not detected:
            return 'utf-8'
        
        # 处理别名
        encoding_map = {
            'GB2312': 'gb2312',
            'GB18030': 'gb18030',
            'ISO-8859-1': 'latin-1',
            'ascii': 'utf-8',  # ASCII 兼容 UTF-8
        }
        return encoding_map.get(detected, detected.lower())
    
    async def extract_text(self, file_path: str, file_type: str) -> Tuple[str, str]:
        """
        从文件中提取文本
        返回: (文本内容, 编码)
        
        根据文件类型调用不同的提取器
        """
        file_type = file_type.lower()
        
        if file_type == 'txt':
            return await self._extract_from_txt(file_path)
        elif file_type == 'epub':
            return await self._extract_from_epub(file_path)
        elif file_type in ['doc', 'docx']:
            return await self._extract_from_word(file_path)
        elif file_type == 'pdf':
            return await self._extract_from_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
    
    async def _extract_from_txt(self, file_path: str) -> Tuple[str, str]:
        """提取 TXT 文件"""
        import aiofiles
        
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()
        
        encoding = self.detect_encoding(content)
        text = content.decode(encoding, errors='replace')
        
        return text, encoding
    
    async def _extract_from_epub(self, file_path: str) -> Tuple[str, str]:
        """提取 EPUB 文件"""
        from ebooklib import epub
        
        book = epub.read_epub(file_path)
        text_parts = []
        
        for item in book.get_items():
            if item.get_type() == 9:  # ITEM_DOCUMENT
                content = item.get_content()
                # 从 HTML 中提取文本
                text = self._html_to_text(content)
                text_parts.append(text)
        
        full_text = '\n\n'.join(text_parts)
        return full_text, 'utf-8'
    
    async def _extract_from_word(self, file_path: str) -> Tuple[str, str]:
        """提取 Word 文件"""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        full_text = '\n'.join(text_parts)
        return full_text, 'utf-8'
    
    async def _extract_from_pdf(self, file_path: str) -> Tuple[str, str]:
        """提取 PDF 文件"""
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = '\n'.join(text_parts)
        return full_text, 'utf-8'
    
    def _html_to_text(self, html_content: bytes) -> str:
        """将 HTML 内容转换为纯文本"""
        import html
        
        # 简单的正则清理
        text = html_content.decode('utf-8', errors='replace')
        
        # 移除 script 和 style
        text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        
        # 处理换行标签
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</p>', '\n\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)
        
        # 移除所有 HTML 标签
        text = re.sub(r'<[^>]+>', '', text)
        
        # 解码 HTML 实体
        text = html.unescape(text)
        
        # 清理多余空白
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        return text.strip()
